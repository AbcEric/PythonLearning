# 读取docx中的文本代码示例
import docx
import win32com.client

'''
# 获取文档对象
file = docx.Document("worddemo.docx")
print("段落数:" + str(len(file.paragraphs)))  # 段落数为13，每个回车隔离一段

# 输出每一段的内容
for para in file.paragraphs:
    print(para.text)

# 输出段落编号及段落内容
for i in range(len(file.paragraphs)):
    print("第" + str(i) + "段的内容是：" + file.paragraphs[i].text)



o = win32com.client.Dispatch("wps.application")
o.Visible=True
doc = o.Documents.Add()
doc.Content.text="Hello world!"
'''
import win32com.client
#新建WPS进程
#wps、et、wpp对应的是金山文件、表格和演示
#word、excel、powerpoint对应的是微软的文字、表格和演示
wpsApp=win32com.client.Dispatch("et.Application")
#设置界面可见
wpsApp.Visible=1
#新建一个wps工作簿
xlBook = wpsApp.Workbooks.Add()
#选定工作簿中活动工作表的某个单元格
cell = xlBook.ActiveSheet.Cells(1,1)
#设置单元格的值
cell.Value='one'
#保存工作簿
xlBook.SaveAs(r"c:/HelloWorld.xls")
xlBook.Close()
wpsApp.Quit()
del wpsApp